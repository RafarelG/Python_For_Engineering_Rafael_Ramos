import os
from faster_whisper import WhisperModel
from docx import Document
from pydub import AudioSegment

# ========= CONFIGURACIÓN =========
AUDIO_FILE = "audio.mp3"   # cambia al nombre de tu archivo
MODEL_SIZE = "small"          # tiny, base, small, medium, large-v3
LANG = "es"
CON_TIMESTAMPS = True
# =================================

# 1. Normalizar audio a 16kHz mono
wav16 = "temp_16k.wav"
audio = AudioSegment.from_file(AUDIO_FILE)
audio = audio.set_channels(1).set_frame_rate(16000)
audio.export(wav16, format="wav")

# 2. Cargar modelo (usa GPU si está disponible)
DEVICE = "cuda" if os.system("nvidia-smi") == 0 else "cpu"
model = WhisperModel(MODEL_SIZE, device=DEVICE, compute_type="int8_float16" if DEVICE=="cuda" else "int8")

# 3. Transcribir
segments, info = model.transcribe(
    wav16,
    language=LANG,
    vad_filter=True,
    vad_parameters=dict(min_silence_duration_ms=500),
    beam_size=1,
    word_timestamps=False,
    temperature=0.0
)

# 4. Guardar en Word
def fmt_ts(t):
    h = int(t//3600); m = int((t%3600)//60); s = int(t%60)
    return f"{h:02d}:{m:02d}:{s:02d}"

doc = Document()
doc.add_heading("Transcripción de Nota de Voz", level=1)

for seg in segments:
    if CON_TIMESTAMPS:
        doc.add_paragraph(f"[{fmt_ts(seg.start)}–{fmt_ts(seg.end)}] {seg.text.strip()}")
    else:
        doc.add_paragraph(seg.text.strip())

doc.save("transcripcion.docx")
print("✅ Transcripción guardada en transcripcion.docx")
